#!/usr/bin/env python3
"""
Extract text from Word document preserving structure
"""
from docx import Document
import sys

def extract_docx_content(filepath):
    """Extract content from docx file"""
    doc = Document(filepath)
    
    output = []
    output.append(f"=" * 80)
    output.append(f"DOCUMENT ANALYSIS: {filepath}")
    output.append(f"=" * 80)
    output.append(f"Total Paragraphs: {len(doc.paragraphs)}")
    output.append(f"Total Sections: {len(doc.sections)}")
    output.append(f"Total Tables: {len(doc.tables)}")
    output.append(f"=" * 80)
    output.append("")
    
    current_page = 1
    para_count = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        para_count += 1
        
        # Detect chapter headings
        if text.startswith("CHAPTER") or text.startswith("Chapter"):
            output.append(f"\n{'=' * 80}")
            output.append(f"PARAGRAPH {para_count} [Index {i}]")
            output.append(f"{'=' * 80}")
            output.append(text)
            output.append(f"{'=' * 80}\n")
        # Detect section headings (numbered sections)
        elif len(text) < 100 and (text[0].isdigit() or text.startswith("5.")):
            output.append(f"\n{'-' * 60}")
            output.append(f"[Para {para_count}] HEADING: {text}")
            output.append(f"{'-' * 60}")
        else:
            # Regular paragraph
            output.append(f"[{para_count}] {text}")
    
    # Extract tables
    if doc.tables:
        output.append(f"\n\n{'=' * 80}")
        output.append(f"TABLES FOUND: {len(doc.tables)}")
        output.append(f"{'=' * 80}")
        for i, table in enumerate(doc.tables):
            output.append(f"\nTable {i+1}:")
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                output.append(f"  {row_text}")
    
    return "\n".join(output)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <path_to_docx>")
        sys.exit(1)
    
    content = extract_docx_content(sys.argv[1])
    print(content)
